#!/usr/bin/env python3
"""将PPT演讲准备文档从Markdown转为Word格式"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 读取MD文件
with open(r"C:\Users\Lenovo\Desktop\Enoch\演讲准备\PocketVibe_Presentation_演讲准备文档.md", "r", encoding="utf-8") as f:
    lines = f.readlines()

i = 0
in_code_block = False
code_lines = []

def add_rich_paragraph(doc, text, style_name=None):
    """添加支持粗体标记的段落"""
    p = doc.add_paragraph(style=style_name)
    # 分割粗体标记
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_table_from_lines(doc, table_lines):
    """从markdown表格行创建Word表格"""
    rows_data = []
    for line in table_lines:
        line = line.strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        rows_data.append(cells)
    # 去掉分隔行（包含---的行）
    rows_data = [r for r in rows_data if not all(re.match(r'^-+$', c) for c in r)]
    if not rows_data:
        return
    ncols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=ncols, style='Table Grid')
    for ri, row in enumerate(rows_data):
        for ci, cell_text in enumerate(row):
            if ci < ncols:
                table.rows[ri].cells[ci].text = cell_text
                # 表头加粗
                if ri == 0:
                    for paragraph in table.rows[ri].cells[ci].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    # 代码块
    if stripped.startswith('```'):
        if in_code_block:
            # 结束代码块
            code_text = ''.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text.rstrip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 设置段落底纹
            pPr = p.paragraph_format
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            in_code_block = False
            code_lines = []
        else:
            in_code_block = True
            code_lines = []
        i += 1
        continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # 空行
    if not stripped:
        i += 1
        continue
    
    # 水平线
    if stripped == '---':
        doc.add_paragraph('─' * 60)
        i += 1
        continue
    
    # 分隔符行 ═══
    if stripped.startswith('# ═'):
        p = doc.add_paragraph()
        run = p.add_run(stripped.lstrip('# '))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        i += 1
        continue
    
    # 表格
    if stripped.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_from_lines(doc, table_lines)
        continue
    
    # 标题
    if stripped.startswith('#'):
        match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if match:
            level = len(match.group(1))
            title_text = match.group(2)
            if level == 1:
                doc.add_heading(title_text, level=1)
            elif level == 2:
                doc.add_heading(title_text, level=2)
            elif level == 3:
                doc.add_heading(title_text, level=3)
            else:
                doc.add_heading(title_text, level=min(level, 4))
        i += 1
        continue
    
    # 引用块
    if stripped.startswith('>'):
        text = stripped.lstrip('> ')
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1
        continue
    
    # 普通段落（支持粗体）
    add_rich_paragraph(doc, stripped)
    i += 1

# 保存
output_path = r"C:\Users\Lenovo\Desktop\Enoch\演讲准备\PocketVibe_Presentation_演讲准备文档_v2.docx"
doc.save(output_path)
print(f"✅ Word文档已生成: {output_path}")
